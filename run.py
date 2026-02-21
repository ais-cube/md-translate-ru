#!/usr/bin/env python3
"""
run.py - Единый пайплайн: документ на входе -> перевод -> все форматы на выходе.

Принимает: .md, .docx, .pdf (один или папка)
Спрашивает: язык перевода, бюджет, форматы
Выдает: переведенный документ в PDF + DOCX + HTML + MD (все в одном файле)

Требования:
    pip install anthropic rich fpdf2 markdown python-docx pdfplumber

Использование:
    python run.py                          # интерактивный режим
    python run.py --input docs/            # папка с файлами
    python run.py --input report.pdf       # один файл
    python run.py --lang en-de             # английский -> немецкий
    python run.py --dry-run                # только оценить стоимость
    python run.py --no-interactive         # для скриптов/CI

Переменные окружения:
    ANTHROPIC_API_KEY  - ключ API (обязателен для перевода)
    TRANSLATE_MODEL    - модель (по умолчанию: claude-sonnet-4-5-20250929)
"""

import argparse
import json
import os
import re
import signal
import sys
import time
from pathlib import Path
from datetime import datetime

# ---------------------------------------------------------------------------
# Config - сохраняется между запусками
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent
CONFIG_PATH = ROOT / ".run_config.json"

def load_config() -> dict:
    defaults = {"ui_lang": None, "last_lang_pair": "en-ru", "last_output_dir": str(ROOT / "output")}
    if CONFIG_PATH.exists():
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                saved = json.load(f)
            defaults.update(saved)
        except Exception:
            pass
    return defaults

def save_config(cfg: dict):
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

_config = load_config()

# ---------------------------------------------------------------------------
# i18n - двуязычный интерфейс (ru / en)
# ---------------------------------------------------------------------------
UI_STRINGS = {
    # === Общие ===
    "app_title": {
        "ru": "md-translate-ru - единый пайплайн перевода документов",
        "en": "md-translate-ru - unified document translation pipeline",
    },
    "app_subtitle": {
        "ru": "Вход: .md, .docx, .pdf -> Перевод -> Выход: MD + HTML + PDF + DOCX",
        "en": "Input: .md, .docx, .pdf -> Translate -> Output: MD + HTML + PDF + DOCX",
    },
    "choose_ui_lang": {
        "ru": "Язык интерфейса",
        "en": "Interface language",
    },
    "choice": {
        "ru": "Выбор",
        "en": "Choice",
    },

    # === Проверка зависимостей ===
    "checking_deps": {
        "ru": "Проверка зависимостей...",
        "en": "Checking dependencies...",
    },
    "deps_ok": {
        "ru": "Все зависимости установлены",
        "en": "All dependencies installed",
    },
    "deps_missing": {
        "ru": "Не установлены пакеты:",
        "en": "Missing packages:",
    },
    "deps_install_hint": {
        "ru": "Установите командой:",
        "en": "Install with:",
    },

    # === Шаги меню ===
    "step_files": {
        "ru": "Шаг 1: Входные файлы",
        "en": "Step 1: Input files",
    },
    "step_lang": {
        "ru": "Шаг 2: Язык перевода",
        "en": "Step 2: Translation language",
    },
    "step_mode": {
        "ru": "Шаг 3: Режим",
        "en": "Step 3: Mode",
    },
    "step_formats": {
        "ru": "Шаг 4: Выходные форматы",
        "en": "Step 4: Output formats",
    },

    # === Файлы ===
    "files_found": {
        "ru": "Найдены файлы:",
        "en": "Files found:",
    },
    "files_not_found": {
        "ru": "Файлы не найдены автоматически.",
        "en": "No files found automatically.",
    },
    "enter_path": {
        "ru": "Путь к файлу или папке",
        "en": "Path to file or folder",
    },
    "enter_path_manual": {
        "ru": "Ввести путь вручную",
        "en": "Enter path manually",
    },
    "select_files": {
        "ru": "Выберите файлы (номера через запятую, или 'all')",
        "en": "Select files (numbers separated by comma, or 'all')",
    },
    "files_selected": {
        "ru": "Выбрано файлов",
        "en": "Files selected",
    },
    "no_files": {
        "ru": "Нет файлов для обработки",
        "en": "No files to process",
    },
    "path_not_found": {
        "ru": "Путь не найден",
        "en": "Path not found",
    },

    # === Язык ===
    "default_marker": {
        "ru": "(по умолчанию)",
        "en": "(default)",
    },

    # === Режим ===
    "mode_sync": {
        "ru": "Перевести сейчас (синхронно)",
        "en": "Translate now (synchronous)",
    },
    "mode_dry": {
        "ru": "Только оценить стоимость (dry-run)",
        "en": "Estimate cost only (dry-run)",
    },
    "set_budget": {
        "ru": "Задать лимит бюджета?",
        "en": "Set budget limit?",
    },
    "max_budget": {
        "ru": "Максимальный бюджет (USD)",
        "en": "Maximum budget (USD)",
    },

    # === Форматы ===
    "fmt_all": {
        "ru": "Все (MD + HTML + PDF + DOCX)",
        "en": "All (MD + HTML + PDF + DOCX)",
    },
    "fmt_pdf_docx": {
        "ru": "PDF + DOCX",
        "en": "PDF + DOCX",
    },
    "fmt_pdf_only": {
        "ru": "Только PDF",
        "en": "PDF only",
    },
    "fmt_docx_only": {
        "ru": "Только DOCX",
        "en": "DOCX only",
    },
    "fmt_pdf_html": {
        "ru": "PDF + HTML",
        "en": "PDF + HTML",
    },
    "fmt_md_only": {
        "ru": "Только MD",
        "en": "MD only",
    },

    # === Выход ===
    "output_name": {
        "ru": "Имя выходного файла (без расширения)",
        "en": "Output filename (no extension)",
    },
    "output_dir": {
        "ru": "Выходная папка",
        "en": "Output folder",
    },

    # === Пайплайн ===
    "pipeline_title": {
        "ru": "ПАЙПЛАЙН ПЕРЕВОДА",
        "en": "TRANSLATION PIPELINE",
    },
    "step_reading": {
        "ru": "ШАГ 1/4: Чтение входных файлов...",
        "en": "STEP 1/4: Reading input files...",
    },
    "step_forecast": {
        "ru": "ШАГ 2/4: Прогноз стоимости...",
        "en": "STEP 2/4: Cost forecast...",
    },
    "step_translate": {
        "ru": "ШАГ 3/4: Перевод...",
        "en": "STEP 3/4: Translating...",
    },
    "step_generate": {
        "ru": "ШАГ 4/4: Сборка и генерация файлов...",
        "en": "STEP 4/4: Assembly and file generation...",
    },
    "chars": {
        "ru": "символов",
        "en": "characters",
    },
    "files_label": {
        "ru": "файлов",
        "en": "files",
    },
    "total": {
        "ru": "Итого",
        "en": "Total",
    },
    "glossary_terms": {
        "ru": "Глоссарий: {n} терминов",
        "en": "Glossary: {n} terms",
    },
    "empty_file": {
        "ru": "Пустой файл или ошибка чтения",
        "en": "Empty file or read error",
    },
    "no_text": {
        "ru": "Нет текста для перевода",
        "en": "No text to translate",
    },
    "dry_run_note": {
        "ru": "Это была оценка. Уберите --dry-run для запуска перевода.",
        "en": "This was an estimate. Remove --dry-run to start translation.",
    },
    "confirm_start": {
        "ru": "Запустить перевод?",
        "en": "Start translation?",
    },
    "cancelled": {
        "ru": "Отменено.",
        "en": "Cancelled.",
    },
    "direction": {
        "ru": "Направление",
        "en": "Direction",
    },
    "output_formats": {
        "ru": "Форматы выхода",
        "en": "Output formats",
    },
    "ctrlc_hint": {
        "ru": "Ctrl+C - остановить после текущего файла",
        "en": "Ctrl+C - stop after current file",
    },
    "model_label": {
        "ru": "Модель",
        "en": "Model",
    },
    "api_key_missing": {
        "ru": "ОШИБКА: установите ANTHROPIC_API_KEY",
        "en": "ERROR: set ANTHROPIC_API_KEY",
    },
    "budget_exhausted": {
        "ru": "Бюджет исчерпан: ${spent:.2f} / ${budget:.2f}",
        "en": "Budget exhausted: ${spent:.2f} / ${budget:.2f}",
    },
    "stopped_by_user": {
        "ru": "Остановлено пользователем",
        "en": "Stopped by user",
    },
    "no_translated": {
        "ru": "Нет переведенных текстов",
        "en": "No translated texts",
    },
    "fonts_missing": {
        "ru": "Шрифты не найдены в fonts/ - PDF будет пропущен",
        "en": "Fonts not found in fonts/ - PDF will be skipped",
    },

    # === Итоги ===
    "done_title": {
        "ru": "ГОТОВО",
        "en": "DONE",
    },
    "result_title": {
        "ru": "Результат",
        "en": "Result",
    },
    "files_translated": {
        "ru": "Файлов переведено",
        "en": "Files translated",
    },
    "chars_processed": {
        "ru": "Символов",
        "en": "Characters",
    },
    "tokens_label": {
        "ru": "Токены",
        "en": "Tokens",
    },
    "cost_label": {
        "ru": "Стоимость",
        "en": "Cost",
    },
    "time_label": {
        "ru": "Время",
        "en": "Time",
    },
    "output_folder": {
        "ru": "Выходная папка",
        "en": "Output folder",
    },
    "error_label": {
        "ru": "Ошибка",
        "en": "Error",
    },

    # === Прогноз ===
    "forecast_title": {
        "ru": "Прогноз",
        "en": "Forecast",
    },
    "file_col": {
        "ru": "Файл",
        "en": "File",
    },
    "chars_col": {
        "ru": "Символы",
        "en": "Chars",
    },
    "chunks_col": {
        "ru": "Чанки",
        "en": "Chunks",
    },
    "time_col": {
        "ru": "Время",
        "en": "Time",
    },
    "price_col": {
        "ru": "Цена",
        "en": "Price",
    },

    # === Interrupt ===
    "force_exit": {
        "ru": "Принудительный выход.",
        "en": "Force exit.",
    },
    "ctrlc_stop": {
        "ru": "Ctrl+C - перевод остановится после текущего файла. Нажмите еще раз для немедленного выхода.",
        "en": "Ctrl+C - translation will stop after current file. Press again for immediate exit.",
    },

    # === Время ===
    "sec": {"ru": "сек", "en": "sec"},
    "min": {"ru": "мин", "en": "min"},
    "hour": {"ru": "ч", "en": "h"},

    # === Картинки ===
    "images_found": {
        "ru": "Обнаружены изображения в тексте",
        "en": "Images found in text",
    },
    "images_count": {
        "ru": "{n} изображений в {f} файлах",
        "en": "{n} images in {f} files",
    },
    "translate_images_q": {
        "ru": "Переводить alt-текст изображений?",
        "en": "Translate image alt-text?",
    },
}

# Current UI language
_ui_lang = _config.get("ui_lang", "ru")

def t(key: str, **kwargs) -> str:
    """Get translated UI string."""
    entry = UI_STRINGS.get(key, {})
    text = entry.get(_ui_lang, entry.get("ru", key))
    if kwargs:
        try:
            text = text.format(**kwargs)
        except (KeyError, ValueError):
            pass
    return text

def set_ui_lang(lang: str):
    global _ui_lang
    _ui_lang = lang
    _config["ui_lang"] = lang
    save_config(_config)

# ---------------------------------------------------------------------------
# Rich / Fallback
# ---------------------------------------------------------------------------
try:
    from rich.console import Console
    from rich.table import Table
    from rich.panel import Panel
    from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn, TimeElapsedColumn, TimeRemainingColumn
    from rich.prompt import Prompt, Confirm, IntPrompt
    from rich.rule import Rule
    from rich import box
    HAS_RICH = True
except ImportError:
    HAS_RICH = False

console = Console() if HAS_RICH else None

# ---------------------------------------------------------------------------
# Dependency check
# ---------------------------------------------------------------------------
REQUIRED_PACKAGES = {
    "anthropic": "anthropic>=0.40.0",
    "pdfplumber": "pdfplumber>=0.10.0",
    "docx": "python-docx>=1.0.0",
    "fpdf": "fpdf2>=2.8.0",
    "markdown": "markdown>=3.5.0",
}

def check_dependencies() -> list[str]:
    """Check required packages. Returns list of missing package names."""
    missing = []
    for module, pip_name in REQUIRED_PACKAGES.items():
        try:
            __import__(module)
        except ImportError:
            missing.append(pip_name)
    return missing

def run_dependency_check():
    """Check and report missing dependencies."""
    missing = check_dependencies()
    if not missing:
        return True

    print(f"\n  {'ОШИБКА' if _ui_lang == 'ru' else 'ERROR'}: {t('deps_missing')}")
    for pkg in missing:
        print(f"    - {pkg}")
    print(f"\n  {t('deps_install_hint')}")
    print(f"    pip install {' '.join(missing)}")
    print(f"\n  {'Или все сразу:' if _ui_lang == 'ru' else 'Or install all at once:'}")
    print(f"    pip install -r requirements.txt\n")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
FONT_DIR = ROOT / "fonts"
GLOSSARY_PATH = ROOT / "glossary.json"
GLOSSARY_CANDIDATES_PATH = ROOT / "glossary_candidates.json"
TRANSLATE_SPEC = ROOT / "TRANSLATE.md"
HUMANIZER_SPEC = ROOT / "HUMANIZER.md"

# ---------------------------------------------------------------------------
# Settings
# ---------------------------------------------------------------------------
DEFAULT_MODEL = "claude-sonnet-4-5-20250929"
MAX_OUTPUT_TOKENS = 16384
CHUNK_SIZE_CHARS = 40000

# Стоимость (USD за 1M токенов) - Sonnet 4.5
COST_INPUT_PER_M = 3.0
COST_OUTPUT_PER_M = 15.0
AVG_SECONDS_PER_1K_CHARS = 3.5

# Языковые пары
LANGUAGES = {
    "en-ru": ("английского", "русский", "English", "Russian"),
    "ru-en": ("русского", "английский", "Russian", "English"),
    "en-de": ("английского", "немецкий", "English", "German"),
    "en-es": ("английского", "испанский", "English", "Spanish"),
    "en-fr": ("английского", "французский", "English", "French"),
    "en-zh": ("английского", "китайский", "English", "Chinese"),
    "en-ja": ("английского", "японский", "English", "Japanese"),
    "en-pt": ("английского", "португальский", "English", "Portuguese"),
    "de-en": ("немецкого", "английский", "German", "English"),
    "fr-en": ("французского", "английский", "French", "English"),
    "es-en": ("испанского", "английский", "Spanish", "English"),
}

# Поддерживаемые форматы ввода
INPUT_EXTENSIONS = {'.md', '.markdown', '.txt', '.docx', '.doc', '.pdf'}

# ---------------------------------------------------------------------------
# Graceful interrupt
# ---------------------------------------------------------------------------
_interrupted = False

def _signal_handler(signum, frame):
    global _interrupted
    if _interrupted:
        ui_print(f"\n[bold red]{t('force_exit')}[/]" if HAS_RICH else f"\n{t('force_exit')}")
        sys.exit(1)
    _interrupted = True
    ui_print(
        f"\n[yellow]{t('ctrlc_stop')}[/]"
        if HAS_RICH else f"\n{t('ctrlc_stop')}"
    )

signal.signal(signal.SIGINT, _signal_handler)

# ---------------------------------------------------------------------------
# UI Helpers
# ---------------------------------------------------------------------------

def ui_print(msg: str, **kwargs):
    if HAS_RICH:
        console.print(msg, **kwargs)
    else:
        clean = re.sub(r'\[/?[^\]]*\]', '', str(msg))
        print(clean)


def log(msg: str, level: str = "INFO"):
    ts = datetime.now().strftime("%H:%M:%S")
    if HAS_RICH:
        colors = {"INFO": "cyan", "WARN": "yellow", "ERROR": "red", "OK": "green",
                  "STEP": "magenta"}
        c = colors.get(level, "white")
        console.print(f"[dim]{ts}[/] [{c}]{level:>5}[/]  {msg}")
    else:
        print(f"[{ts}] [{level}] {msg}")


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def load_json(path: Path) -> list:
    if not path.exists():
        return []
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except (json.JSONDecodeError, Exception):
        return []


def estimate_tokens(text: str) -> int:
    return len(text) // 3


def calc_cost(input_tokens: int, output_tokens: int) -> float:
    return (input_tokens * COST_INPUT_PER_M + output_tokens * COST_OUTPUT_PER_M) / 1_000_000


def format_duration(seconds: float) -> str:
    if seconds < 60:
        return f"{seconds:.0f} {t('sec')}"
    elif seconds < 3600:
        m, s = divmod(int(seconds), 60)
        return f"{m} {t('min')} {s} {t('sec')}"
    else:
        h, rem = divmod(int(seconds), 3600)
        m = rem // 60
        return f"{h} {t('hour')} {m} {t('min')}"


def format_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} Б"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} КБ"
    else:
        return f"{size_bytes / (1024*1024):.1f} МБ"


# ---------------------------------------------------------------------------
# STEP 1: Input - чтение файлов любого формата
# ---------------------------------------------------------------------------

def extract_text_from_pdf(path: Path) -> str:
    """Извлечь текст из PDF."""
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
        return "\n\n".join(texts)
    except ImportError:
        log("Установите pdfplumber: pip install pdfplumber", "ERROR")
        return ""
    except Exception as e:
        log(f"Ошибка чтения PDF {path.name}: {e}", "ERROR")
        return ""


def extract_text_from_docx(path: Path) -> str:
    """Извлечь текст из DOCX."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Сохранить стиль заголовков
                style = para.style.name if para.style else ""
                if "Heading 1" in style:
                    paragraphs.append(f"# {text}")
                elif "Heading 2" in style:
                    paragraphs.append(f"## {text}")
                elif "Heading 3" in style:
                    paragraphs.append(f"### {text}")
                elif "Heading 4" in style:
                    paragraphs.append(f"#### {text}")
                elif "List" in style:
                    paragraphs.append(f"- {text}")
                else:
                    paragraphs.append(text)

        # Таблицы
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
                rows.insert(1, header_sep)
                paragraphs.append("\n".join(rows))

        return "\n\n".join(paragraphs)
    except ImportError:
        log("Установите python-docx: pip install python-docx", "ERROR")
        return ""
    except Exception as e:
        log(f"Ошибка чтения DOCX {path.name}: {e}", "ERROR")
        return ""


def read_input_file(path: Path) -> tuple[str, str]:
    """Прочитать файл любого формата. Возвращает (text, original_format)."""
    ext = path.suffix.lower()

    if ext in ('.md', '.markdown', '.txt'):
        text = path.read_text(encoding="utf-8")
        return text, "md"

    elif ext == '.pdf':
        log(f"  Извлечение текста из PDF: {path.name}...")
        text = extract_text_from_pdf(path)
        return text, "pdf"

    elif ext in ('.docx', '.doc'):
        if ext == '.doc':
            log(f"  .doc формат - попытка чтения как .docx: {path.name}", "WARN")
        log(f"  Извлечение текста из DOCX: {path.name}...")
        text = extract_text_from_docx(path)
        return text, "docx"

    else:
        log(f"Неподдерживаемый формат: {ext}", "ERROR")
        return "", "unknown"


def detect_images(text: str) -> list[tuple[str, str]]:
    """Find all image references ![alt](path) in Markdown text."""
    return re.findall(r'!\[([^\]]*)\]\(([^)]+)\)', text)


def discover_input_files(input_path: Path) -> list[Path]:
    """Найти все поддерживаемые файлы в указанном пути."""
    if input_path.is_file():
        if input_path.suffix.lower() in INPUT_EXTENSIONS:
            return [input_path]
        else:
            log(f"Неподдерживаемый формат: {input_path.suffix}", "ERROR")
            return []

    if input_path.is_dir():
        files = []
        for ext in INPUT_EXTENSIONS:
            files.extend(input_path.glob(f"*{ext}"))
        return sorted(set(files))

    log(f"Путь не найден: {input_path}", "ERROR")
    return []


# ---------------------------------------------------------------------------
# STEP 2: Translation - перевод через Claude API
# ---------------------------------------------------------------------------

def build_system_prompt(lang_pair: str, glossary: list) -> str:
    """Построить системный промпт для выбранной языковой пары."""
    lang_from, lang_to, _, _ = LANGUAGES.get(lang_pair, LANGUAGES["en-ru"])

    parts = []
    parts.append(f"Ты - профессиональный технический переводчик с {lang_from} на {lang_to}.")
    parts.append("Следуй приведенным ниже правилам ТОЧНО и БЕЗ ОТКЛОНЕНИЙ.\n")

    if TRANSLATE_SPEC.exists():
        spec = TRANSLATE_SPEC.read_text(encoding="utf-8")
        parts.append("=" * 60)
        parts.append("СПЕЦИФИКАЦИЯ ПЕРЕВОДА (TRANSLATE.md)")
        parts.append("=" * 60)
        parts.append(spec)

    if HUMANIZER_SPEC.exists():
        humanizer = HUMANIZER_SPEC.read_text(encoding="utf-8")
        parts.append("\n" + "=" * 60)
        parts.append("РЕДАКТОРСКИЕ ПРАВИЛА (HUMANIZER.md) - только anti-AI cleanup")
        parts.append("ВАЖНО: НЕ применять секцию PERSONALITY AND SOUL.")
        parts.append("НЕ добавлять личный голос, эмоции, юмор, первое лицо.")
        parts.append("Использовать ТОЛЬКО для устранения AI-штампов.")
        parts.append("=" * 60)
        parts.append(humanizer)

    if glossary:
        parts.append("\n" + "=" * 60)
        parts.append("КАНОНИЧЕСКИЙ ГЛОССАРИЙ (glossary.json)")
        parts.append("Если term встречается в тексте - использовать ТОЛЬКО перевод из глоссария.")
        parts.append("=" * 60)
        glossary_text = "\n".join(
            f"- {e.get('term_en', e.get('term', ''))} -> {e.get('term_ru', e.get('translation', ''))}"
            for e in glossary if isinstance(e, dict)
        )
        parts.append(glossary_text)

    return "\n\n".join(parts)


def build_user_prompt(source_text: str, filename: str, lang_pair: str,
                      is_chunk: bool = False, chunk_num: int = 0,
                      total_chunks: int = 0,
                      translate_images: bool = False) -> str:
    """Построить пользовательский промпт для перевода."""
    lang_from, lang_to, lang_from_en, lang_to_en = LANGUAGES.get(lang_pair, LANGUAGES["en-ru"])

    chunk_info = ""
    if is_chunk:
        chunk_info = f"\n\nЭто чанк {chunk_num}/{total_chunks}. Переводи только этот фрагмент."

    image_rule = ("8. Изображения ![alt](path) - ПЕРЕВЕДИ alt-text, путь оставь без изменений."
                  if translate_images else
                  "8. Изображения ![alt](path) - оставь ПОЛНОСТЬЮ как есть, включая alt-text и путь.")

    return f"""Переведи следующий документ с {lang_from} на {lang_to}.

Файл: {filename}{chunk_info}

ПРАВИЛА ФОРМАТИРОВАНИЯ (КРИТИЧЕСКИ ВАЖНО):

1. MARKDOWN-СТРУКТУРА: сохрани 1:1 - заголовки (#), списки (-), таблицы (|), code blocks (```), ссылки, изображения.
2. ЗАГОЛОВКИ: каждый заголовок (#, ##, ###) ровно ОДИН РАЗ, на ОДНОЙ строке. Пример: "## Глава 1: Название".
   - ЗАПРЕЩЕНО дублировать заголовки.
   - ЗАПРЕЩЕНО разбивать заголовок на несколько строк.
3. ТАБЛИЦЫ: сохрани markdown-таблицы как таблицы. Не превращай таблицу в текст. Каждая строка таблицы = строка с |.
4. ПАРАГРАФЫ: один параграф = один непрерывный блок текста. НЕ разбивай предложения на отдельные строки.
   - Между параграфами - одна пустая строка.
   - НЕ добавляй лишних пустых строк.
   - Длинные предложения НЕ разбивай переносами строк - оставляй одним абзацем.
5. CODE BLOCKS: НЕ переводи содержимое ``` блоков. Оставь как есть.
6. URL, идентификаторы, тикеры: НЕ переводи.
7. Глоссарий: используй ТОЛЬКО канонические термины (если предоставлен).
{image_rule}
9. Содержание (Table of Contents): переведи ОДИН РАЗ, не дублируй.
10. Иерархия: # = заголовок документа, ## = разделы, ### = подразделы. Сохрани уровни.

ФОРМАТ ВЫВОДА:
- Верни ТОЛЬКО переведенный Markdown. Без комментариев, без преамбулы, без пост-скриптума.
- Не оборачивай в ```markdown``` блоки.

---НАЧАЛО ИСХОДНОГО ТЕКСТА---

{source_text}

---КОНЕЦ ИСХОДНОГО ТЕКСТА---"""


def split_into_chunks(text: str, max_chars: int = CHUNK_SIZE_CHARS) -> list[str]:
    """Разбить текст на чанки по заголовкам ##."""
    if len(text) <= max_chars:
        return [text]

    sections = re.split(r'(?=\n## )', text)
    chunks = []
    current_chunk = ""

    for section in sections:
        if len(current_chunk) + len(section) <= max_chars:
            current_chunk += section
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            if len(section) > max_chars:
                paragraphs = section.split("\n\n")
                current_chunk = ""
                for para in paragraphs:
                    if len(current_chunk) + len(para) + 2 <= max_chars:
                        current_chunk += ("\n\n" if current_chunk else "") + para
                    else:
                        if current_chunk.strip():
                            chunks.append(current_chunk.strip())
                        current_chunk = para
            else:
                current_chunk = section

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks if chunks else [text]


def translate_text(client, model: str, system_prompt: str,
                   source_text: str, filename: str, lang_pair: str,
                   is_chunk: bool = False, chunk_num: int = 0,
                   total_chunks: int = 0,
                   translate_images: bool = False) -> tuple[str, int, int]:
    """Перевести текст через Claude API."""
    user_prompt = build_user_prompt(source_text, filename, lang_pair,
                                    is_chunk, chunk_num, total_chunks,
                                    translate_images)

    response = client.messages.create(
        model=model,
        max_tokens=MAX_OUTPUT_TOKENS,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}],
    )

    translated = response.content[0].text
    return translated, response.usage.input_tokens, response.usage.output_tokens


def translate_document(client, model: str, system_prompt: str,
                       source_text: str, filename: str, lang_pair: str,
                       translate_images: bool = False) -> dict:
    """Перевести один документ (с разбивкой на чанки)."""
    chunks = split_into_chunks(source_text)

    stats = {
        "file": filename,
        "source_chars": len(source_text),
        "chunks": len(chunks),
        "input_tokens": 0,
        "output_tokens": 0,
        "cost": 0.0,
        "status": "ok",
        "translated_text": "",
    }

    translated_parts = []

    for i, chunk in enumerate(chunks, 1):
        if _interrupted:
            stats["status"] = "interrupted"
            return stats

        is_chunked = len(chunks) > 1
        if is_chunked:
            log(f"    Чанк {i}/{len(chunks)} ({len(chunk):,} символов)...")

        try:
            translation, inp_tok, out_tok = translate_text(
                client, model, system_prompt,
                chunk, filename, lang_pair,
                is_chunk=is_chunked, chunk_num=i, total_chunks=len(chunks),
                translate_images=translate_images
            )
            translated_parts.append(translation)
            stats["input_tokens"] += inp_tok
            stats["output_tokens"] += out_tok
        except Exception as e:
            log(f"    ОШИБКА: {e}", "ERROR")
            stats["status"] = f"error: {e}"
            return stats

        if is_chunked and i < len(chunks):
            time.sleep(1)

    stats["translated_text"] = "\n\n".join(translated_parts)
    stats["cost"] = calc_cost(stats["input_tokens"], stats["output_tokens"])
    return stats


# ---------------------------------------------------------------------------
# STEP 3: Assembly - сборка в один документ
# ---------------------------------------------------------------------------

# Attribution lines by target language
ATTRIBUTION = {
    "ru": "Переведено с помощью **https://github.com/ais-cube/md-translate-ru**",
    "en": "Translated with **https://github.com/ais-cube/md-translate-ru**",
    "de": "Ubersetzt mit **https://github.com/ais-cube/md-translate-ru**",
    "es": "Traducido con **https://github.com/ais-cube/md-translate-ru**",
    "fr": "Traduit avec **https://github.com/ais-cube/md-translate-ru**",
    "zh": "使用 **https://github.com/ais-cube/md-translate-ru** 翻译",
    "ja": "**https://github.com/ais-cube/md-translate-ru** で翻訳",
    "pt": "Traduzido com **https://github.com/ais-cube/md-translate-ru**",
}


def dedup_lines(text: str) -> str:
    """Remove consecutive duplicate lines, split headings, and orphaned fragments."""
    lines = text.split('\n')
    result = []
    prev_stripped = None
    skip_next_if_fragment = None  # text fragment to skip if found on next line

    for line in lines:
        stripped = line.strip()

        # Skip exact consecutive duplicates (non-empty)
        if stripped and stripped == prev_stripped:
            continue

        # Skip orphaned fragment from a split heading
        if skip_next_if_fragment and stripped == skip_next_if_fragment:
            skip_next_if_fragment = None
            continue
        skip_next_if_fragment = None

        # Handle consecutive heading lines: keep longer, mark fragment for skip
        if (prev_stripped and stripped
            and re.match(r'^#{1,6}\s', prev_stripped)
            and re.match(r'^#{1,6}\s', stripped)):
            prev_level = re.match(r'^(#{1,6})\s', prev_stripped).group(1)
            curr_level = re.match(r'^(#{1,6})\s', stripped).group(1)
            prev_text = re.sub(r'^#{1,6}\s+', '', prev_stripped)
            curr_text = re.sub(r'^#{1,6}\s+', '', stripped)

            if prev_level == curr_level:
                if prev_text in curr_text:
                    # Current is fuller version - replace previous
                    result[-1] = line
                    prev_stripped = stripped
                    continue
                elif curr_text in prev_text:
                    # Previous is already fuller - skip this and mark fragment
                    remainder = prev_text.replace(curr_text, '').strip()
                    if remainder:
                        skip_next_if_fragment = remainder
                    continue

        result.append(line)
        prev_stripped = stripped

    return '\n'.join(result)


def get_target_lang_code(lang_pair: str) -> str:
    """Extract target language code from pair like 'en-ru' -> 'ru'."""
    return lang_pair.split('-')[-1] if '-' in lang_pair else "ru"


def cleanup_markdown(text: str) -> str:
    """Clean up translated markdown: fix line breaks, tables, headings."""
    lines = text.split('\n')
    result = []
    in_code = False
    in_table = False
    prev_blank = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # Toggle code block state
        if line.strip().startswith('```'):
            in_code = not in_code
            result.append(line)
            prev_blank = False
            i += 1
            continue

        # Inside code block - keep as is
        if in_code:
            result.append(line)
            prev_blank = False
            i += 1
            continue

        stripped = line.strip()

        # Blank line handling - max one between content
        if not stripped:
            if not prev_blank and result:
                result.append('')
                prev_blank = True
            i += 1
            continue

        prev_blank = False

        # Table lines - keep intact, ensure table continuity
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table and result and result[-1] != '':
                result.append('')
            in_table = True
            result.append(line)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                # Don't add blank line if next content will add its own
                if result and result[-1] != '':
                    result.append('')

        # Heading - must be single line, ensure blank line before
        if re.match(r'^#{1,6}\s', stripped):
            if result and result[-1] != '':
                result.append('')
            result.append(line)
            i += 1
            continue

        # Horizontal rules
        if stripped in ('---', '***', '___'):
            result.append(line)
            i += 1
            continue

        # List items - keep as separate lines
        if re.match(r'^(\s*)([-*+]|\d+\.)\s+', line):
            result.append(line)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            result.append(line)
            i += 1
            continue

        # Image reference
        if re.match(r'^!\[', stripped):
            result.append(line)
            i += 1
            continue

        # Regular paragraph text - join consecutive non-special lines
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            next_stripped = lines[j].strip()
            # Stop joining at: blank line, heading, list, table, code, blockquote, hr, image
            if (not next_stripped
                or re.match(r'^#{1,6}\s', next_stripped)
                or re.match(r'^(\s*)([-*+]|\d+\.)\s+', lines[j])
                or (next_stripped.startswith('|') and next_stripped.endswith('|'))
                or next_stripped.startswith('```')
                or next_stripped.startswith('>')
                or next_stripped in ('---', '***', '___')
                or re.match(r'^!\[', next_stripped)):
                break
            para_lines.append(next_stripped)
            j += 1

        # Join paragraph into single line
        paragraph = ' '.join(para_lines)
        result.append(paragraph)
        i = j
        continue

    # Final cleanup: remove trailing blank lines
    while result and result[-1] == '':
        result.pop()

    return '\n'.join(result) + '\n'


def assemble_document(translations: list[dict], title: str = "",
                      lang_pair: str = "en-ru") -> str:
    """Собрать все переводы в один Markdown-документ."""
    target_lang = get_target_lang_code(lang_pair)
    attribution = ATTRIBUTION.get(target_lang, ATTRIBUTION["en"])

    parts = []

    # Attribution at the top
    parts.append(f"> {attribution}\n")

    for i, tr in enumerate(translations):
        if len(translations) > 1 and i > 0:
            parts.append(f"\n---\n")
        parts.append(tr["translated_text"])

    assembled = "\n\n".join(parts)

    # Deduplicate consecutive identical headings/lines
    assembled = dedup_lines(assembled)

    # Clean up markdown formatting
    assembled = cleanup_markdown(assembled)

    return assembled


# ---------------------------------------------------------------------------
# STEP 4: Output - генерация всех форматов
# ---------------------------------------------------------------------------

def generate_html(md_text: str, title: str, lang_pair: str = "en-ru") -> str:
    """Markdown -> HTML with proper formatting."""
    import markdown
    extensions = [
        'markdown.extensions.tables',
        'markdown.extensions.fenced_code',
        'markdown.extensions.toc',
        'markdown.extensions.sane_lists',
        'markdown.extensions.nl2br',
        'markdown.extensions.smarty',
    ]
    html_content = markdown.markdown(md_text, extensions=extensions)
    date_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    target_lang = get_target_lang_code(lang_pair)

    # Map short code to HTML lang attribute
    lang_html = {
        "ru": "ru", "en": "en", "de": "de", "es": "es",
        "fr": "fr", "zh": "zh", "ja": "ja", "pt": "pt",
    }.get(target_lang, "en")

    return f"""<!DOCTYPE html>
<html lang="{lang_html}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{ box-sizing: border-box; }}
        body {{
            font-family: 'Segoe UI', 'DejaVu Sans', 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.75; color: #1a1a2e; background: #fafbfc;
            max-width: 860px; margin: 0 auto; padding: 40px 32px;
            word-wrap: break-word; overflow-wrap: break-word;
        }}
        p {{
            margin: 0 0 1em 0;
            text-align: justify;
            hyphens: auto;
            -webkit-hyphens: auto;
            -ms-hyphens: auto;
        }}
        h1 {{
            font-size: 2em; font-weight: 700;
            border-bottom: 3px solid #4361ee;
            padding-bottom: 12px; margin: 40px 0 20px 0;
            line-height: 1.3;
        }}
        h2 {{
            font-size: 1.5em; font-weight: 700;
            border-bottom: 1px solid #dee2e6;
            padding-bottom: 8px; margin: 36px 0 16px 0;
            line-height: 1.3;
        }}
        h3 {{ font-size: 1.25em; font-weight: 700; margin: 28px 0 12px 0; line-height: 1.3; }}
        h4 {{ font-size: 1.1em; font-weight: 700; margin: 24px 0 10px 0; }}
        a {{ color: #4361ee; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        code {{
            background: #e9ecef; padding: 2px 6px; border-radius: 4px;
            font-family: 'Consolas', 'Courier New', monospace;
            font-size: 0.9em; color: #d63384;
            word-break: break-all;
        }}
        pre {{
            background: #1e1e2e; color: #cdd6f4; padding: 20px;
            border-radius: 8px; overflow-x: auto; line-height: 1.5;
            margin: 16px 0; white-space: pre-wrap; word-wrap: break-word;
        }}
        pre code {{ background: none; color: inherit; padding: 0; word-break: normal; }}
        ul, ol {{ padding-left: 28px; margin: 8px 0 16px 0; }}
        li {{ margin: 4px 0; line-height: 1.6; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
        th, td {{ border: 1px solid #dee2e6; padding: 10px 14px; text-align: left; }}
        th {{ background: #f1f3f5; font-weight: 600; }}
        tr:nth-child(even) {{ background: #f8f9fa; }}
        blockquote {{
            border-left: 4px solid #4361ee; margin: 16px 0;
            padding: 12px 20px; background: #eef2ff;
            color: #495057; font-style: italic;
        }}
        blockquote p {{ margin: 0 0 0.5em 0; }}
        blockquote p:last-child {{ margin: 0; }}
        img {{ max-width: 100%; height: auto; margin: 16px 0; display: block; }}
        hr {{ border: none; border-top: 2px solid #dee2e6; margin: 32px 0; }}
        .meta {{
            color: #868e96; font-size: 0.85em;
            border-top: 1px solid #dee2e6;
            padding-top: 16px; margin-top: 48px;
        }}
        @media print {{
            body {{ max-width: none; padding: 20px; }}
            pre {{ white-space: pre-wrap; }}
            a {{ color: #1a1a2e; text-decoration: underline; }}
        }}
        @media (max-width: 640px) {{
            body {{ padding: 16px; }}
            h1 {{ font-size: 1.6em; }}
            h2 {{ font-size: 1.3em; }}
            table {{ font-size: 0.9em; }}
            th, td {{ padding: 6px 8px; }}
        }}
    </style>
</head>
<body>
{html_content}
<div class="meta">md-translate-ru | {date_str}</div>
</body>
</html>"""


def generate_pdf(md_text: str, output_path: Path, title: str):
    """Markdown -> PDF через fpdf2 с кириллицей."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Регистрация шрифтов
    font_map = {
        "DejaVuSans": {"": "DejaVuSans.ttf", "B": "DejaVuSans-Bold.ttf",
                        "I": "DejaVuSans-Oblique.ttf", "BI": "DejaVuSans-BoldOblique.ttf"},
        "DejaVuMono": {"": "DejaVuSansMono.ttf", "B": "DejaVuSansMono-Bold.ttf"},
    }
    for family, styles in font_map.items():
        for style, filename in styles.items():
            font_path = FONT_DIR / filename
            if font_path.exists():
                pdf.add_font(family, style, str(font_path))

    pdf.add_page()
    pdf.set_font("DejaVuSans", size=10)

    def clean(text):
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = re.sub(r'!\[(.+?)\]\(.+?\)', r'[\1]', text)
        return text

    lines = md_text.split('\n')
    in_code = False
    code_buf = []

    for line in lines:
        if line.strip().startswith('```'):
            if in_code:
                pdf.set_fill_color(30, 30, 46)
                pdf.set_text_color(205, 214, 244)
                pdf.set_font("DejaVuMono", size=8)
                for cl in code_buf:
                    t = cl[:95] + "..." if len(cl) > 95 else cl
                    pdf.cell(0, 4.5, t)
                    pdf.ln(4.5)
                pdf.set_font("DejaVuSans", size=10)
                pdf.set_text_color(30, 30, 30)
                pdf.ln(4)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buf.append(line)
            continue

        stripped = line.strip()

        if not stripped:
            pdf.ln(3)
            continue

        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            sizes = {1: 20, 2: 16, 3: 13, 4: 11}
            pdf.ln(4)
            pdf.set_font("DejaVuSans", "B", sizes.get(level, 10))
            pdf.set_text_color(26, 26, 46)
            pdf.multi_cell(0, sizes.get(level, 10) * 0.6, clean(text))
            if level <= 2:
                pdf.set_draw_color(67, 97, 238)
                pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + pdf.epw, pdf.get_y())
            pdf.ln(3)
            pdf.set_font("DejaVuSans", size=10)
            pdf.set_text_color(30, 30, 30)
            continue

        if stripped in ('---', '***', '___'):
            pdf.ln(4)
            pdf.set_draw_color(200, 200, 200)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + pdf.epw, pdf.get_y())
            pdf.ln(8)
            continue

        if stripped.startswith('|') and stripped.endswith('|'):
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            n = len(cells) or 1
            w = pdf.epw / n
            pdf.set_font("DejaVuSans", size=9)
            for c in cells:
                t = clean(c)[:40]
                pdf.cell(w, 7, t, border=1)
            pdf.ln()
            pdf.set_font("DejaVuSans", size=10)
            continue

        list_m = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)', line)
        if list_m:
            indent = len(list_m.group(1))
            offset = 10 + (indent // 2) * 6
            pdf.set_x(pdf.l_margin + offset)
            marker = list_m.group(2)
            bullet = "\u2022 " if marker in ('-', '*', '+') else f"{marker} "
            pdf.multi_cell(pdf.epw - offset - 4, 6, bullet + clean(list_m.group(3)))
            pdf.ln(1)
            continue

        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            pdf.set_x(pdf.l_margin + 8)
            pdf.set_font("DejaVuSans", "I", 10)
            pdf.set_text_color(59, 59, 92)
            pdf.multi_cell(pdf.epw - 16, 6, clean(text))
            pdf.ln(2)
            pdf.set_font("DejaVuSans", size=10)
            pdf.set_text_color(30, 30, 30)
            continue

        pdf.multi_cell(0, 6, clean(stripped))
        pdf.ln(2)

    pdf.output(str(output_path))


def generate_docx(md_text: str, output_path: Path, title: str):
    """Markdown -> DOCX через python-docx."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor

    doc = DocxDocument()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    lines = md_text.split('\n')
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        rows_data = []
        for row_line in table_buf:
            if re.match(r'^\|[\s\-:|]+\|$', row_line.strip()):
                continue
            cells = [c.strip() for c in row_line.strip().split('|')[1:-1]]
            if cells:
                rows_data.append(cells)

        if rows_data:
            n_cols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=n_cols)
            table.style = 'Table Grid'
            for i, row_data in enumerate(rows_data):
                for j, cell_text in enumerate(row_data):
                    if j < n_cols:
                        table.rows[i].cells[j].text = clean_md(cell_text)
            if rows_data:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
        table_buf = []

    def clean_md(text):
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = re.sub(r'!\[(.+?)\]\(.+?\)', r'[\1]', text)
        return text

    for line in lines:
        if line.strip().startswith('```'):
            if in_code:
                flush_table()
                code_text = '\n'.join(code_buf)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x20, 0x20, 0x30)
                p.paragraph_format.left_indent = Inches(0.3)
                code_buf = []
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_buf.append(line)
            continue

        stripped = line.strip()

        if stripped.startswith('|') and stripped.endswith('|'):
            table_buf.append(line)
            continue
        else:
            flush_table()

        if not stripped:
            continue

        if stripped.startswith('#'):
            level = min(len(stripped) - len(stripped.lstrip('#')), 4)
            text = stripped.lstrip('#').strip()
            doc.add_heading(clean_md(text), level=level)
            continue

        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            continue

        list_m = re.match(r'^(\s*)([-*+])\s+(.+)', line)
        if list_m:
            doc.add_paragraph(clean_md(list_m.group(3)), style='List Bullet')
            continue

        num_m = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if num_m:
            doc.add_paragraph(clean_md(num_m.group(2)), style='List Number')
            continue

        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(clean_md(text))
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x3B, 0x3B, 0x5C)
            continue

        doc.add_paragraph(clean_md(stripped))

    flush_table()
    doc.save(str(output_path))


def generate_outputs(md_text: str, output_dir: Path, base_name: str,
                     title: str, formats: list[str],
                     lang_pair: str = "en-ru") -> dict:
    """Сгенерировать все выходные форматы."""
    results = {"files": [], "errors": []}
    output_dir.mkdir(parents=True, exist_ok=True)

    if "md" in formats:
        try:
            path = output_dir / f"{base_name}.md"
            path.write_text(md_text, encoding="utf-8")
            results["files"].append(("MD", path))
            log(f"  MD:   {path.name}", "OK")
        except Exception as e:
            results["errors"].append(f"MD: {e}")

    if "html" in formats:
        try:
            path = output_dir / f"{base_name}.html"
            html = generate_html(md_text, title, lang_pair)
            path.write_text(html, encoding="utf-8")
            results["files"].append(("HTML", path))
            log(f"  HTML: {path.name}", "OK")
        except Exception as e:
            results["errors"].append(f"HTML: {e}")

    if "pdf" in formats:
        try:
            path = output_dir / f"{base_name}.pdf"
            generate_pdf(md_text, path, title)
            results["files"].append(("PDF", path))
            log(f"  PDF:  {path.name}", "OK")
        except Exception as e:
            results["errors"].append(f"PDF: {e}")

    if "docx" in formats:
        try:
            path = output_dir / f"{base_name}.docx"
            generate_docx(md_text, path, title)
            results["files"].append(("DOCX", path))
            log(f"  DOCX: {path.name}", "OK")
        except Exception as e:
            results["errors"].append(f"DOCX: {e}")

    return results


# ---------------------------------------------------------------------------
# Interactive Menu
# ---------------------------------------------------------------------------

def find_input_candidates() -> list[tuple[Path, int, str]]:
    """Найти файлы с поддерживаемыми форматами."""
    candidates = []
    skip = {'.git', '__pycache__', 'fonts', 'output', 'node_modules', '.venv', 'venv'}

    def scan(base: Path, depth: int = 0):
        if depth > 2:
            return
        try:
            for p in sorted(base.iterdir()):
                if p.name.startswith('.') or p.name in skip:
                    continue
                if p.is_file() and p.suffix.lower() in INPUT_EXTENSIONS:
                    candidates.append((p, p.stat().st_size, p.suffix.lower()))
                elif p.is_dir():
                    scan(p, depth + 1)
        except PermissionError:
            pass

    scan(ROOT)
    cwd = Path.cwd()
    if cwd.resolve() != ROOT.resolve():
        scan(cwd)

    seen = set()
    unique = []
    for item in candidates:
        resolved = item[0].resolve()
        if resolved not in seen:
            seen.add(resolved)
            unique.append(item)
    return unique


def interactive_menu() -> dict:
    """Полное интерактивное меню с выбором языка интерфейса."""
    global _ui_lang

    result = {
        "input_files": [],
        "lang_pair": _config.get("last_lang_pair", "en-ru"),
        "budget": None,
        "formats": ["md", "html", "pdf", "docx"],
        "output_dir": Path(_config.get("last_output_dir", str(ROOT / "output"))),
        "output_name": "translated",
        "dry_run": False,
        "batch": False,
        "model": None,
    }

    if HAS_RICH:
        # ========== 0. ЯЗЫК ИНТЕРФЕЙСА ==========
        if _config.get("ui_lang") is None:
            console.print()
            console.print("  [bold cyan] 1[/]  Русский")
            console.print("  [bold cyan] 2[/]  English")
            console.print()
            lang_ch = Prompt.ask(t("choose_ui_lang") + " / Interface language",
                                  choices=["1", "2"], default="1")
            set_ui_lang("ru" if lang_ch == "1" else "en")
        else:
            console.print()
            saved = "Русский" if _ui_lang == "ru" else "English"
            console.print(f"  [dim]{t('choose_ui_lang')}: {saved}  (s - switch)[/]")
            switch = Prompt.ask(t("choice"), default="")
            if switch.strip().lower() == "s":
                set_ui_lang("en" if _ui_lang == "ru" else "ru")
                console.print(f"  [green]{'English' if _ui_lang == 'en' else 'Русский'}[/]")

        console.print()
        console.print(Panel(
            f"[bold cyan]md-translate-ru[/] - {t('app_title')}\n"
            f"[dim]{t('app_subtitle')}[/]",
            title="run.py",
            border_style="cyan",
        ))

        # ========== 1. ВХОДНЫЕ ФАЙЛЫ ==========
        console.print()
        console.print(Rule(f"[bold]{t('step_files')}[/]", style="cyan"))
        console.print()

        files_found = find_input_candidates()

        if files_found:
            console.print(f"[bold]{t('files_found')}[/]\n")
            for i, (f, size, ext) in enumerate(files_found, 1):
                try:
                    rel = f.relative_to(ROOT)
                except ValueError:
                    rel = f
                ext_colors = {'.md': 'green', '.pdf': 'red', '.docx': 'blue', '.txt': 'yellow'}
                c = ext_colors.get(ext, 'white')
                console.print(f"  [bold cyan]{i:2d}[/]  [{c}]{ext:6s}[/]  {rel}  [dim]{format_size(size)}[/]")

            console.print()
            console.print(f"  [bold cyan] 0[/]  {t('enter_path_manual')}")
            console.print()

            choice = Prompt.ask(t("select_files"), default="all")

            if choice.strip() == "0":
                path_str = Prompt.ask(t("enter_path"))
                p = Path(path_str)
                if p.exists():
                    result["input_files"] = discover_input_files(p)
                else:
                    log(f"{t('path_not_found')}: {path_str}", "ERROR")
                    sys.exit(1)
            elif choice.strip().lower() == "all":
                result["input_files"] = [f[0] for f in files_found]
            else:
                indices = []
                for part in choice.split(","):
                    part = part.strip()
                    if "-" in part:
                        a, b = part.split("-", 1)
                        indices.extend(range(int(a), int(b) + 1))
                    elif part.isdigit():
                        indices.append(int(part))
                for idx in indices:
                    if 1 <= idx <= len(files_found):
                        result["input_files"].append(files_found[idx - 1][0])
        else:
            console.print(f"[yellow]{t('files_not_found')}[/]")
            path_str = Prompt.ask(t("enter_path"))
            p = Path(path_str)
            if p.exists():
                result["input_files"] = discover_input_files(p)
            else:
                log(f"{t('path_not_found')}: {path_str}", "ERROR")
                sys.exit(1)

        if not result["input_files"]:
            log(t("no_files"), "ERROR")
            sys.exit(1)

        console.print(f"\n  [green]{t('files_selected')}: {len(result['input_files'])}[/]")

        # ========== 2. ЯЗЫК ПЕРЕВОДА ==========
        console.print()
        console.print(Rule(f"[bold]{t('step_lang')}[/]", style="cyan"))
        console.print()

        lang_options = list(LANGUAGES.keys())
        default_idx = "1"
        for i, lp in enumerate(lang_options, 1):
            _, _, from_en, to_en = LANGUAGES[lp]
            marker = f" [bold green]({t('default_marker')})[/]" if lp == result["lang_pair"] else ""
            if lp == result["lang_pair"]:
                default_idx = str(i)
            console.print(f"  [bold cyan]{i:2d}[/]  {from_en} -> {to_en}  [dim]({lp})[/]{marker}")

        console.print()
        lang_choice = Prompt.ask(t("choice"), default=default_idx)
        try:
            idx = int(lang_choice) - 1
            if 0 <= idx < len(lang_options):
                result["lang_pair"] = lang_options[idx]
        except ValueError:
            if lang_choice in LANGUAGES:
                result["lang_pair"] = lang_choice

        _config["last_lang_pair"] = result["lang_pair"]
        save_config(_config)

        _, _, from_en, to_en = LANGUAGES[result["lang_pair"]]
        console.print(f"  [green]{from_en} -> {to_en}[/]")

        # ========== 3. РЕЖИМ ==========
        console.print()
        console.print(Rule(f"[bold]{t('step_mode')}[/]", style="cyan"))
        console.print()

        modes = [
            ("1", t("mode_sync"), "sync"),
            ("2", t("mode_dry"), "dry"),
        ]
        for key, label, _ in modes:
            console.print(f"  [bold cyan]{key}[/]  {label}")
        console.print()

        mode_choice = Prompt.ask(t("choice"), choices=["1", "2"], default="1")
        mode = next(m[2] for m in modes if m[0] == mode_choice)
        result["dry_run"] = mode == "dry"

        # ========== 4. БЮДЖЕТ ==========
        if not result["dry_run"]:
            console.print()
            if Confirm.ask(t("set_budget"), default=False):
                budget_str = Prompt.ask(t("max_budget"), default="10.00")
                try:
                    result["budget"] = float(budget_str)
                except ValueError:
                    pass

        # ========== 5. ВЫХОДНЫЕ ФОРМАТЫ ==========
        console.print()
        console.print(Rule(f"[bold]{t('step_formats')}[/]", style="cyan"))
        console.print()

        fmt_options = [
            ("1", t("fmt_all"), ["md", "html", "pdf", "docx"]),
            ("2", t("fmt_pdf_docx"), ["pdf", "docx"]),
            ("3", t("fmt_pdf_only"), ["pdf"]),
            ("4", t("fmt_docx_only"), ["docx"]),
            ("5", t("fmt_pdf_html"), ["pdf", "html"]),
            ("6", t("fmt_md_only"), ["md"]),
        ]
        for key, label, _ in fmt_options:
            console.print(f"  [bold cyan]{key}[/]  {label}")
        console.print()

        fmt_choice = Prompt.ask(t("choice"), choices=[f[0] for f in fmt_options], default="1")
        result["formats"] = next(f[2] for f in fmt_options if f[0] == fmt_choice)

        # ========== 6. ВЫХОДНОЕ ИМЯ ==========
        console.print()
        default_name = "translated"
        if len(result["input_files"]) == 1:
            default_name = result["input_files"][0].stem + "_translated"
        result["output_name"] = Prompt.ask(t("output_name"), default=default_name)

        # ========== 7. ВЫХОДНАЯ ПАПКА ==========
        default_out = str(result["output_dir"])
        out_str = Prompt.ask(t("output_dir"), default=default_out)
        result["output_dir"] = Path(out_str).resolve()
        _config["last_output_dir"] = str(result["output_dir"])
        save_config(_config)

    else:
        # === Fallback без Rich ===
        print(f"\n{'='*55}")
        # Language selection
        if _config.get("ui_lang") is None:
            print("  1. Русский")
            print("  2. English")
            lang_ch = input("Language [1]: ").strip() or "1"
            set_ui_lang("ru" if lang_ch == "1" else "en")

        print(t("app_title"))
        print(f"{'='*55}")

        files_found = find_input_candidates()
        if files_found:
            print(f"\n{t('files_found')}")
            for i, (f, size, ext) in enumerate(files_found, 1):
                print(f"  {i:2d}. [{ext}] {f.name} ({format_size(size)})")
            print(f"   0. {t('enter_path_manual')}")
            choice = input(f"\n{t('select_files')} [all]: ").strip() or "all"
            if choice == "0":
                p = Path(input(f"{t('enter_path')}: ").strip())
                result["input_files"] = discover_input_files(p) if p.exists() else []
            elif choice.lower() == "all":
                result["input_files"] = [f[0] for f in files_found]
            else:
                for part in choice.split(","):
                    idx = int(part.strip())
                    if 1 <= idx <= len(files_found):
                        result["input_files"].append(files_found[idx - 1][0])
        else:
            p = Path(input(f"{t('enter_path')}: ").strip())
            result["input_files"] = discover_input_files(p) if p.exists() else []

        if not result["input_files"]:
            print(t("no_files")); sys.exit(1)

        print(f"\n{t('step_lang')}:")
        print("  1=EN->RU 2=RU->EN 3=EN->DE 4=EN->ES 5=EN->FR")
        lang_map = {"1": "en-ru", "2": "ru-en", "3": "en-de", "4": "en-es", "5": "en-fr"}
        result["lang_pair"] = lang_map.get(input(f"{t('choice')} [1]: ").strip() or "1", "en-ru")

        print(f"\n{t('step_mode')}:")
        print(f"  1={t('mode_sync')} 2={t('mode_dry')}")
        mode = input(f"{t('choice')} [1]: ").strip() or "1"
        result["dry_run"] = mode == "2"

        b = input(f"{t('max_budget')} (Enter=no limit): ").strip()
        if b:
            try: result["budget"] = float(b)
            except ValueError: pass

        print(f"\n{t('step_formats')}:")
        print(f"  1={t('fmt_all')} 2={t('fmt_pdf_docx')} 3={t('fmt_pdf_only')} 4={t('fmt_docx_only')} 5={t('fmt_md_only')}")
        fmt_map = {"1": ["md","html","pdf","docx"], "2": ["pdf","docx"],
                   "3": ["pdf"], "4": ["docx"], "5": ["md"]}
        result["formats"] = fmt_map.get(input(f"{t('choice')} [1]: ").strip() or "1",
                                         ["md","html","pdf","docx"])

        default_name = "translated"
        if len(result["input_files"]) == 1:
            default_name = result["input_files"][0].stem + "_translated"
        result["output_name"] = input(f"{t('output_name')} [{default_name}]: ").strip() or default_name

        result["output_dir"] = Path(
            input(f"{t('output_dir')} [{ROOT / 'output'}]: ").strip() or str(ROOT / "output")
        ).resolve()

    return result


# ---------------------------------------------------------------------------
# Forecast
# ---------------------------------------------------------------------------

def show_forecast(input_files: list[Path], texts: list[str], system_tokens: int,
                  budget: float = None) -> list[dict]:
    """Показать прогноз стоимости."""
    forecasts = []
    for f, text in zip(input_files, texts):
        chunks = split_into_chunks(text)
        est_input = estimate_tokens(text) + system_tokens
        est_output = int(estimate_tokens(text) * 1.15)
        est_cost = calc_cost(est_input, est_output)
        est_time = len(text) / 1000 * AVG_SECONDS_PER_1K_CHARS
        forecasts.append({
            "name": f.name, "chars": len(text), "chunks": len(chunks),
            "est_cost": est_cost, "est_time": est_time,
        })

    total_cost = sum(fc["est_cost"] for fc in forecasts)
    total_time = sum(fc["est_time"] for fc in forecasts)
    total_chars = sum(fc["chars"] for fc in forecasts)

    if HAS_RICH:
        table = Table(title=t("forecast_title"), box=box.ROUNDED, show_lines=True, title_style="bold cyan")
        table.add_column("#", style="dim", width=4, justify="right")
        table.add_column(t("file_col"), style="bold white", max_width=30)
        table.add_column(t("chars_col"), justify="right", style="cyan")
        table.add_column(t("chunks_col"), justify="center")
        table.add_column(t("time_col"), justify="right", style="yellow")
        table.add_column(t("price_col"), justify="right", style="green")

        for i, fc in enumerate(forecasts, 1):
            table.add_row(str(i), fc["name"], f"{fc['chars']:,}",
                          str(fc["chunks"]), format_duration(fc["est_time"]),
                          f"${fc['est_cost']:.2f}")

        table.add_section()
        budget_note = ""
        if budget is not None:
            r = budget - total_cost
            budget_lbl = "budget" if _ui_lang == "en" else "бюджет"
            remain_lbl = "remaining" if _ui_lang == "en" else "остаток"
            short_lbl = "short" if _ui_lang == "en" else "не хватает"
            budget_note = (f"  [green]{budget_lbl} ${budget:.2f}, {remain_lbl} ${r:.2f}[/]" if r >= 0
                           else f"  [red]{budget_lbl} ${budget:.2f}, {short_lbl} ${-r:.2f}[/]")

        table.add_row("", f"[bold]{t('total').upper()}: {len(forecasts)} {t('files_label')}[/]", f"[bold]{total_chars:,}[/]",
                       "", f"[bold]{format_duration(total_time)}[/]",
                       f"[bold]${total_cost:.2f}[/]" + budget_note)
        console.print()
        console.print(table)
        console.print()
    else:
        print(f"\n{t('forecast_title').upper()}: {len(forecasts)} {t('files_label')}, {total_chars:,} {t('chars')}")
        print(f"{t('time_col')}: {format_duration(total_time)}, {t('price_col')}: ${total_cost:.2f}")
        if budget is not None:
            print(f"Budget: ${budget:.2f}")
        print()

    return forecasts


# ---------------------------------------------------------------------------
# Main Pipeline
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Unified pipeline: document -> translate -> all formats",
    )
    parser.add_argument("--input", help="File or folder to translate")
    parser.add_argument("--lang", default=None, help="Language pair (en-ru, en-de, ...)")
    parser.add_argument("--budget", type=float, default=None, help="Budget limit USD")
    parser.add_argument("--format", default=None, help="Formats: all, pdf, docx, html, md")
    parser.add_argument("--output", default=None, help="Output filename")
    parser.add_argument("--output-dir", default=None, help="Output folder")
    parser.add_argument("--dry-run", action="store_true", help="Estimate cost only")
    parser.add_argument("--model", default=None, help="Claude model")
    parser.add_argument("--no-interactive", action="store_true", help="No interactive menu")
    parser.add_argument("--ui-lang", default=None, choices=["ru", "en"], help="Interface language")
    args = parser.parse_args()

    # === Проверка зависимостей (после --help) ===
    run_dependency_check()

    if args.ui_lang:
        set_ui_lang(args.ui_lang)

    is_interactive = (
        not args.no_interactive
        and args.input is None
        and sys.stdin.isatty()
    )

    if is_interactive:
        config = interactive_menu()
    else:
        if not args.input:
            log(t("no_files"), "ERROR")
            sys.exit(1)

        input_path = Path(args.input)
        input_files = discover_input_files(input_path)
        if not input_files:
            sys.exit(1)

        fmt_map = {
            "all": ["md", "html", "pdf", "docx"],
            "pdf": ["pdf"], "docx": ["docx"], "html": ["html"], "md": ["md"],
        }
        formats = fmt_map.get(args.format or "all", ["md", "html", "pdf", "docx"])

        default_name = "translated"
        if len(input_files) == 1:
            default_name = input_files[0].stem + "_translated"

        config = {
            "input_files": input_files,
            "lang_pair": args.lang or "en-ru",
            "budget": args.budget,
            "formats": formats,
            "output_dir": Path(args.output_dir or str(ROOT / "output")).resolve(),
            "output_name": args.output or default_name,
            "dry_run": args.dry_run,
            "model": args.model,
        }

    input_files = config["input_files"]
    lang_pair = config["lang_pair"]
    _, _, from_en, to_en = LANGUAGES.get(lang_pair, LANGUAGES["en-ru"])

    # ========================
    # PIPELINE START
    # ========================

    if HAS_RICH:
        console.print()
        console.print(Rule(f"[bold magenta]{t('pipeline_title')}[/]", style="magenta"))

    # ----- STEP 1 -----
    log(t("step_reading"), "STEP")

    source_texts = []
    for f in input_files:
        log(f"  {f.name} ({f.suffix.lower()})...")
        text, fmt = read_input_file(f)
        if text.strip():
            source_texts.append(text)
            log(f"    {len(text):,} {t('chars')}", "OK")
        else:
            log(f"    {t('empty_file')}", "WARN")

    if not source_texts:
        log(t("no_text"), "ERROR")
        return

    total_chars = sum(len(tx) for tx in source_texts)
    log(f"  {t('total')}: {len(source_texts)} {t('files_label')}, {total_chars:,} {t('chars')}")

    # ----- IMAGE DETECTION -----
    translate_alt_text = False
    total_images = 0
    files_with_images = 0
    for txt in source_texts:
        imgs = detect_images(txt)
        if imgs:
            total_images += len(imgs)
            files_with_images += 1

    if total_images > 0:
        log(f"  {t('images_found')}: {t('images_count', n=total_images, f=files_with_images)}")
        if sys.stdin.isatty():
            if HAS_RICH:
                translate_alt_text = Confirm.ask(f"  {t('translate_images_q')}", default=False)
            else:
                ans = input(f"  {t('translate_images_q')} [y/N]: ").strip().lower()
                translate_alt_text = ans in ("y", "yes", "д", "да")

    if total_images > 0:
        if not translate_alt_text:
            log(f"    -> alt-text: keep original")
        else:
            log(f"    -> alt-text: translate")

    # ----- STEP 2 -----
    log(t("step_forecast"), "STEP")

    glossary = load_json(GLOSSARY_PATH)
    system_prompt = build_system_prompt(lang_pair, glossary)

    system_tokens = estimate_tokens(system_prompt)

    if glossary:
        log(f"  {t('glossary_terms', n=len(glossary))}")

    forecasts = show_forecast(input_files, source_texts, system_tokens, config["budget"])

    if config["dry_run"]:
        ui_print(f"[dim]{t('dry_run_note')}[/]" if HAS_RICH else t("dry_run_note"))
        return

    total_cost = sum(fc["est_cost"] for fc in forecasts)
    if HAS_RICH:
        console.print(f"  {t('direction')}: [bold]{from_en} -> {to_en}[/]")
        console.print(f"  {t('output_formats')}: [bold]{', '.join(f.upper() for f in config['formats'])}[/]")
        console.print(f"  [dim]{t('ctrlc_hint')}[/]")
        if not Confirm.ask(f"\n{t('confirm_start')}", default=True):
            ui_print(f"[dim]{t('cancelled')}[/]")
            return
    else:
        answer = input(f"{t('confirm_start')} [Y/n]: ").strip().lower()
        if answer not in ("", "y", "yes", "д", "да"):
            return

    # ----- STEP 3 -----
    log(t("step_translate"), "STEP")

    model = config.get("model") or os.getenv("TRANSLATE_MODEL", DEFAULT_MODEL)
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        log(t("api_key_missing"), "ERROR")
        log("  export ANTHROPIC_API_KEY='sk-ant-...'")
        sys.exit(1)

    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    log(f"  {t('model_label')}: {model}")
    log(f"  {t('direction')}: {from_en} -> {to_en}")

    translations = []
    total_stats = {"input_tokens": 0, "output_tokens": 0, "cost": 0.0, "errors": 0}
    spent = 0.0
    start_time = time.time()

    for i, (f, text) in enumerate(zip(input_files, source_texts), 1):
        if _interrupted:
            log(t("stopped_by_user"), "WARN")
            break

        if config["budget"] is not None:
            fc = forecasts[i - 1]
            if spent + fc["est_cost"] > config["budget"]:
                log(t("budget_exhausted", spent=spent, budget=config["budget"]), "WARN")
                break

        if HAS_RICH:
            console.print(Rule(f"[bold]{f.name}[/]", style="cyan"))

        log(f"  [{i}/{len(input_files)}] {f.name} ({len(text):,} {t('chars')})...")

        stats = translate_document(client, model, system_prompt, text, f.name, lang_pair,
                                   translate_images=translate_alt_text)

        if stats["translated_text"]:
            translations.append(stats)
            total_stats["input_tokens"] += stats["input_tokens"]
            total_stats["output_tokens"] += stats["output_tokens"]
            spent += stats["cost"]
            log(f"    {stats['input_tokens']:,} in + {stats['output_tokens']:,} out = ${stats['cost']:.2f}", "OK")
        else:
            total_stats["errors"] += 1
            log(f"    {t('error_label')}: {stats['status']}", "ERROR")

        if i < len(input_files):
            time.sleep(2)

    if not translations:
        log(t("no_translated"), "ERROR")
        return

    elapsed = time.time() - start_time

    # ----- STEP 4 -----
    log(t("step_generate"), "STEP")

    assembled_md = assemble_document(translations, lang_pair=lang_pair)

    output_dir = config["output_dir"]
    base_name = config["output_name"]

    title_match = re.search(r'^#\s+(.+)', assembled_md, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else base_name

    if "pdf" in config["formats"]:
        if not (FONT_DIR / "DejaVuSans.ttf").exists():
            log(t("fonts_missing"), "WARN")
            config["formats"] = [f for f in config["formats"] if f != "pdf"]

    gen_results = generate_outputs(assembled_md, output_dir, base_name,
                                    title, config["formats"], lang_pair)

    # ========================
    # RESULTS
    # ========================
    total_cost_actual = calc_cost(total_stats["input_tokens"], total_stats["output_tokens"])

    if HAS_RICH:
        console.print()
        console.print(Rule(f"[bold green]{t('done_title')}[/]", style="green"))
        console.print()

        summary = Table(box=box.SIMPLE, show_header=False, padding=(0, 2))
        summary.add_column("Key", style="dim")
        summary.add_column("Value", style="bold")

        summary.add_row(t("files_translated"), f"{len(translations)} / {len(input_files)}")
        summary.add_row(t("direction"), f"{from_en} -> {to_en}")
        summary.add_row(t("chars_processed"), f"{sum(tx['source_chars'] for tx in translations):,}")
        summary.add_row(t("tokens_label"), f"{total_stats['input_tokens']:,} in + {total_stats['output_tokens']:,} out")
        summary.add_row(t("cost_label"), f"${total_cost_actual:.2f}")
        summary.add_row(t("time_label"), format_duration(elapsed))
        summary.add_row("", "")
        summary.add_row(t("output_folder"), str(output_dir))
        for fmt, path in gen_results["files"]:
            summary.add_row(f"  {fmt}", path.name)

        if gen_results["errors"]:
            for err in gen_results["errors"]:
                summary.add_row(f"[red]{t('error_label')}[/]", err)

        console.print(Panel(summary, title=t("result_title"), border_style="green"))
    else:
        print(f"\n{'='*55}")
        print(f"{t('done_title')}: {len(translations)} {t('files_label')}, ${total_cost_actual:.2f}, {format_duration(elapsed)}")
        print(f"{t('output_folder')}: {output_dir}/")
        for fmt, path in gen_results["files"]:
            print(f"  {fmt}: {path.name}")
        print()


if __name__ == "__main__":
    main()
